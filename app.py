import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Problemas y soluciones de la Escuela Austríaca de Economía", page_icon="📚", layout="wide")

# Function to set the background color
def set_background_color(color):
    st.markdown(
        f"""
        <style>
        .stApp {{
            background-color: {color};
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    ### Sobre esta aplicación

    Esta aplicación proporciona soluciones basadas en la Escuela Austríaca de Economía a diversos problemas económicos y sociales, adaptadas a países de América Latina. Permite a los usuarios obtener soluciones creativas y propuestas de la economía austriaca para diferentes problemas.

    ### Cómo usar la aplicación:

    1. Elija un problema económico/social de la lista predefinida o proponga su propio problema.
    2. Seleccione un país de América Latina.
    3. Haga clic en "Obtener solución" para generar las respuestas.
    4. Lea las soluciones y fuentes proporcionadas.
    5. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, 26 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Problemas y soluciones de la Escuela Austríaca de Economía* [Aplicación web]. https://solucionesau.econom.streamlit.app
    """)

# Titles and Main Column
st.title("Problemas y soluciones de la Escuela Austríaca de Economía")

# Set background color to light yellow
set_background_color("#FFF9C4")  # Light yellow color code

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPER_API_KEY = st.secrets["SERPER_API_KEY"]

    # List of 101 economic and social problems
    problemas_economicos = sorted([
        "¿Cómo reducir el desempleo?", "¿Cómo fomentar el crecimiento económico?", "¿Cómo controlar la inflación?", 
        "¿Cómo mejorar la distribución de la riqueza?", "¿Cómo reducir la pobreza?", "¿Cómo incentivar la inversión?", 
        "¿Cómo fomentar la innovación?", "¿Cómo mejorar la educación?", "¿Cómo proporcionar atención médica asequible?", 
        "¿Cómo reducir la deuda pública?", "¿Cómo mejorar la infraestructura?", "¿Cómo fomentar el comercio internacional?", 
        "¿Cómo regular los monopolios?", "¿Cómo promover la competencia?", "¿Cómo gestionar una crisis financiera?", 
        "¿Cómo aumentar la productividad laboral?", "¿Cómo mejorar el bienestar social?", "¿Cómo combatir el cambio climático?", 
        "¿Cómo promover la energía sostenible?", "¿Cómo incentivar el transporte público?", "¿Cómo asegurar una vivienda asequible?", 
        "¿Cómo reducir los impuestos sin afectar los servicios públicos?", "¿Cómo mejorar la seguridad social?", 
        "¿Cómo fomentar el desarrollo rural?", "¿Cómo mejorar el acceso a la tecnología?", "¿Cómo apoyar a las pequeñas empresas?", 
        "¿Cómo regular las criptomonedas?", "¿Cómo enfrentar el envejecimiento de la población?", "¿Cómo mejorar el sistema de pensiones?", 
        "¿Cómo reducir el déficit comercial?", "¿Cómo manejar la migración?", "¿Cómo reducir la corrupción?", 
        "¿Cómo fortalecer las instituciones democráticas?", "¿Cómo mejorar la seguridad jurídica?", 
        "¿Cómo proteger los derechos de propiedad?", "¿Cómo fomentar el desarrollo industrial?", 
        "¿Cómo diversificar la economía?", "¿Cómo mejorar la eficiencia del mercado laboral?", "¿Cómo reducir la desigualdad salarial?", 
        "¿Cómo promover el desarrollo sostenible?", "¿Cómo aumentar la transparencia en el gobierno?", 
        "¿Cómo mejorar la movilidad social?", "¿Cómo fomentar el respeto a los derechos humanos?", 
        "¿Cómo asegurar la estabilidad financiera?", "¿Cómo promover el ahorro y la inversión?", 
        "¿Cómo mejorar la regulación financiera?", "¿Cómo combatir la evasión fiscal?", 
        "¿Cómo procurar la justicia económica?", "¿Cómo fomentar el desarrollo científico?", 
        "¿Cómo incrementar el acceso a financiamiento para emprendedores?", "¿Cómo mejorar el control de la contaminación?", 
        "¿Cómo fomentar la responsabilidad corporativa?", "¿Cómo asegurar la calidad en la educación?", 
        "¿Cómo fortalecer el sistema de salud?", "¿Cómo reducir el trabajo informal?", 
        "¿Cómo proteger los derechos de los trabajadores?", "¿Cómo asegurar la protección del consumidor?", 
        "¿Cómo fomentar el turismo sostenible?", "¿Cómo reducir la dependencia económica de un sector específico?", 
        "¿Cómo mejorar la eficiencia del gasto público?", "¿Cómo asegurar la estabilidad monetaria?", 
        "¿Cómo reducir la desigualdad de género?", "¿Cómo mejorar las condiciones laborales?", 
        "¿Cómo enfrentar el terrorismo?", "¿Cómo mejorar la seguridad pública?", 
        "¿Cómo fomentar la inversión en investigación y desarrollo?", "¿Cómo diversificar las fuentes de energía?", 
        "¿Cómo fomentar el emprendimiento?", "¿Cómo aumentar la participación ciudadana?", 
        "¿Cómo combatir el tráfico de drogas?", "¿Cómo mejorar la cohesión social?", 
        "¿Cómo fortalecer el liderazgo comunitario?", "¿Cómo asegurar el acceso universal a internet?", 
        "¿Cómo reducir el desperdicio de recursos?", "¿Cómo mejorar la relación entre el sector público y privado?", 
        "¿Cómo enfrentar la automatización y el desempleo tecnológico?", "¿Cómo promover la equidad?", 
        "¿Cómo mejorar la educación financiera?", "¿Cómo fomentar la inclusión social?", 
        "¿Cómo reducir la brecha digital?", "¿Cómo mejorar la calidad del aire?", 
        "¿Cómo gestionar los recursos hídricos?", "¿Cómo promover la salud mental?", 
        "¿Cómo mejorar la resiliencia económica ante desastres naturales?", "¿Cómo fomentar el comercio justo?", 
        "¿Cómo asegurar la protección de datos?", "¿Cómo mejorar el sistema de justicia?", 
        "¿Cómo reducir la impunidad?", "¿Cómo fomentar la cultura del ahorro?", 
        "¿Cómo mejorar la eficiencia energética?", "¿Cómo promover la formación continua?", 
        "¿Cómo asegurar la protección del medio ambiente?", "¿Cómo reducir el endeudamiento de los hogares?", 
        "¿Cómo fomentar la cooperación internacional?", "¿Cómo asegurar el acceso a medicamentos esenciales?", 
        "¿Cómo promover el desarrollo urbano sostenible?", "¿Cómo mejorar la calidad del transporte?", 
        "¿Cómo gestionar los residuos sólidos?", "¿Cómo reducir el uso de plásticos?", 
        "¿Cómo asegurar una transición energética justa?", "¿Cómo promover la participación de mujeres en la economía?"
    ])

    # List of Latin American countries
    paises_latinoamerica = [
        "Argentina", "Bolivia", "Brasil", "Chile", "Colombia", 
        "Costa Rica", "Cuba", "Ecuador", "El Salvador", "Guatemala", 
        "Honduras", "México", "Nicaragua", "Panamá", "Paraguay", 
        "Perú", "República Dominicana", "Uruguay", "Venezuela"
    ]

    def buscar_informacion(query, pais):
        url = "https://google.serper.dev/search"
        payload = json.dumps({
            "q": f"{query} {pais} Escuela Austríaca"
        })
        headers = {
            'X-API-KEY': SERPER_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()

    def generar_respuesta(problema, pais, contexto):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Contexto: {contexto}\n\nProblema: {problema}\nPaís: {pais}\n\nProporciona una solución basada en las propuestas de la Escuela Austríaca de Economía al problema '{problema}' para el país {pais}. La solución debe incluir principios y teorías de la economía austriaca, y si es posible, ejemplos de estudios de caso o estrategias relacionadas.\n\nSolución:",
            "max_tokens": 2048,
            "temperature": 0.6,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["Problema:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(problema, respuestas, fuentes):
        doc = Document()
        doc.add_heading('Diccionario de Problemas y Soluciones Económicas', 0)

        doc.add_heading('Problema', level=1)
        doc.add_paragraph(problema)

        for pais, respuesta in respuestas.items():
            doc.add_heading(f'Solución para el país {pais}', level=2)
            doc.add_paragraph(respuesta)

        doc.add_heading('Fuentes', level=1)

        # Limitar la lista de fuentes a las primeras 10
        for fuente in fuentes[:10]:
            doc.add_paragraph(fuente, style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas y especialistas en la materia para un análisis más profundo.')

        return doc

    st.write("**Elige un problema económico o social de la lista o propón tu propio problema**:")

    opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio problema"])

    if opcion == "Elegir de la lista":
        problema = st.selectbox("Selecciona un problema:", problemas_economicos)
    else:
        problema = st.text_input("Ingresa tu propio problema económico o social:")

    st.write("Selecciona un país de América Latina:")
    pais_seleccionado = st.selectbox("País", paises_latinoamerica)

    if st.button("Obtener solución"):
        if problema and pais_seleccionado:
            with st.spinner("Buscando información y generando soluciones..."):
                respuestas, todas_fuentes = {}, []

                # Buscar información relevante
                resultados_busqueda = buscar_informacion(problema, pais_seleccionado)
                contexto = "\n".join([item["snippet"] for item in resultados_busqueda.get("organic", [])])
                fuentes = [item["link"] for item in resultados_busqueda.get("organic", [])]

                # Generar respuesta basada en la Escuela Austríaca de Economía
                respuesta = generar_respuesta(problema, pais_seleccionado, contexto)
                respuestas[pais_seleccionado] = respuesta
                todas_fuentes.extend(fuentes)

                # Mostrar las respuestas
                st.subheader(f"Soluciones para el problema: {problema}")
                for pais, respuesta in respuestas.items():
                    st.markdown(f"**{pais}:** {respuesta}")

                # Botón para descargar el documento
                doc = create_docx(problema, respuestas, todas_fuentes)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar solución en DOCX",
                    data=buffer,
                    file_name=f"Solución_{problema.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, selecciona un problema y un país.")
